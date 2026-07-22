import csv
import base64
import json
import math
import os
import re
from collections import defaultdict
from datetime import date
from io import BytesIO
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from flask import Blueprint, current_app, render_template_string, request, send_file


zbi_bp = Blueprint("zbi", __name__, url_prefix="/zbi")
BASE_DIR = Path(__file__).resolve().parent
PRODUCTS_FILE = BASE_DIR / "zbi_products.csv"
DELIVERY_FILE = BASE_DIR / "zbi_delivery.csv"
LOGO_B64_FILE = BASE_DIR / "company_logo.b64"
MOSCOW_CENTER_LAT = 55.7558
MOSCOW_CENTER_LON = 37.6176
COMPANY = {
    "full_name": "Общество с ограниченной ответственностью «АР-ФАРВАТЕР»",
    "short_name": "ООО «АР-ФАРВАТЕР»",
    "address": "111677, г. Москва, вн. тер. г. муниципальный округ Некрасовка, ул. Покровская, д. 16, кв. 265",
    "phone": "8-916-727-36-87",
    "email": "nzzk@mail.ru",
    "inn": "9721261781",
    "kpp": "772101001",
    "ogrn": "1267700015802",
    "director": "Никитин Артём Сергеевич",
    "bank": "ООО «Банк Точка»",
    "bik": "044525104",
    "correspondent_account": "30101810745374525104",
    "account": "40702810420000283529",
}


def _float(value, default=0.0):
    try:
        return float(str(value).replace(" ", "").replace(",", "."))
    except (TypeError, ValueError):
        return default


def parse_dimensions(size_text, weight_kg=0):
    """Return transport envelope dimensions in metres and its volume.

    The catalogue contains dimensions written with х, x, × and *.  For products
    without three dimensions we use a conservative volume estimate based on the
    concrete weight and mark it as estimated.
    """
    normalized = str(size_text or "").lower()
    for marker in ("×", "х", "x", "*", "õ"):
        normalized = normalized.replace(marker, "x")
    numbers = [_float(value) for value in re.findall(r"\d+(?:[.,]\d+)?", normalized)]
    numbers = [value for value in numbers if value > 0]
    if len(numbers) >= 3:
        metres = [value / 1000 for value in numbers[:3]]
        length, width, height = sorted(metres, reverse=True)
        return {
            "length_m": length,
            "width_m": width,
            "height_m": height,
            "volume_m3": length * width * height,
            "estimated": False,
        }
    estimated_volume = max(_float(weight_kg) / 2400 * 1.35, 0.03)
    return {
        "length_m": 0,
        "width_m": 0,
        "height_m": 0,
        "volume_m3": estimated_volume,
        "estimated": True,
    }


def load_catalog():
    with PRODUCTS_FILE.open("r", encoding="utf-8-sig", newline="") as handle:
        products = list(csv.DictReader(handle))
    for index, item in enumerate(products):
        if item.get("supplier") == "ИП Михайлов" and str(item.get("name") or "").strip().upper().startswith("ПД "):
            item["group"] = "Лотки и водоотвод"
        item["id"] = str(index)
        item["price_rub"] = _float(item.get("price_rub"))
        item["weight_kg"] = _float(item.get("weight_kg"))
        item["dimensions"] = parse_dimensions(item.get("size_mm"), item["weight_kg"])
    with DELIVERY_FILE.open("r", encoding="utf-8-sig", newline="") as handle:
        delivery = list(csv.DictReader(handle))
    for item in delivery:
        for key in ("capacity_t", "fixed_moscow_rub", "rate_rub_km", "lat", "lon"):
            item[key] = _float(item.get(key))
    return products, delivery


def geocode(address):
    text = str(address or "").strip()
    if not text:
        return None, None, False
    if "," in text:
        parts = [part.strip() for part in text.split(",")]
        if len(parts) == 2:
            lat, lon = _float(parts[0], None), _float(parts[1], None)
            if lat is not None and lon is not None and -90 <= lat <= 90 and -180 <= lon <= 180:
                return lat, lon, is_moscow_point(lat, lon)
    api_key = current_app.config.get("YANDEX_API_KEY") or os.getenv("YANDEX_API_KEY", "")
    if not api_key:
        return None, None, False
    try:
        response = requests.get(
            "https://geocode-maps.yandex.ru/1.x/",
            params={"apikey": api_key, "geocode": text, "format": "json", "lang": "ru_RU"},
            timeout=15,
        )
        response.raise_for_status()
        member = response.json()["response"]["GeoObjectCollection"]["featureMember"][0]
        geo_object = member["GeoObject"]
        lon, lat = geo_object["Point"]["pos"].split()
        components = (
            geo_object.get("metaDataProperty", {})
            .get("GeocoderMetaData", {})
            .get("Address", {})
            .get("Components", [])
        )
        names = {str(component.get("name", "")).strip().casefold() for component in components}
        is_moscow = "москва" in names or "город москва" in names
        return float(lat), float(lon), is_moscow
    except Exception:
        return None, None, False


def haversine_km(lat1, lon1, lat2, lon2):
    radius = 6371.0
    lat1, lon1, lat2, lon2 = [math.radians(value) for value in (lat1, lon1, lat2, lon2)]
    dlat, dlon = lat2 - lat1, lon2 - lon1
    value = math.sin(dlat / 2) ** 2 + math.cos(lat1) * math.cos(lat2) * math.sin(dlon / 2) ** 2
    return radius * 2 * math.asin(math.sqrt(value))


def is_moscow_point(lat, lon):
    """Fallback for coordinates entered without an address.

    The main city is approximated by the MKAD radius; the second area covers
    the principal New Moscow corridor. Normal text addresses use the Yandex
    administrative region returned by the geocoder instead.
    """
    inside_mkad = haversine_km(lat, lon, MOSCOW_CENTER_LAT, MOSCOW_CENTER_LON) <= 24
    inside_new_moscow = 55.20 <= lat <= 55.62 and 36.80 <= lon <= 37.62
    return inside_mkad or inside_new_moscow


def road_distance_km(start_lat, start_lon, end_lat, end_lon):
    coordinates = f"{start_lon},{start_lat};{end_lon},{end_lat}"
    try:
        response = requests.get(
            f"https://router.project-osrm.org/route/v1/driving/{coordinates}",
            params={"overview": "false", "alternatives": "false", "steps": "false"},
            timeout=15,
        )
        data = response.json()
        if data.get("code") == "Ok":
            return data["routes"][0]["distance"] / 1000, "по автодороге"
    except Exception:
        pass
    return haversine_km(start_lat, start_lon, end_lat, end_lon) * 1.25, "оценка по прямой × 1,25"


def vehicle_profile(option):
    name = str(option.get("vehicle") or "").casefold()
    if "манипулятор" in name:
        length, width, height = 6.5, 2.45, 2.5
    else:
        length, width, height = 13.6, 2.45, 2.5
    return {
        "length_m": length,
        "width_m": width,
        "height_m": height,
        "usable_volume_m3": length * width * height * 0.78,
    }


def mixed_delivery_calculation(
    lines, distance, delivery_options, is_moscow, moscow_base_distance, manual_loads=None
):
    manual_loads = manual_loads or {}
    total_weight_kg = sum(line["weight_total_kg"] for line in lines)
    total_volume_m3 = sum(line["volume_total_m3"] for line in lines)
    alternatives = []
    for option in delivery_options:
        capacity_kg = option["capacity_t"] * 1000
        if capacity_kg <= 0:
            continue
        profile = vehicle_profile(option)
        oversize = []
        for line in lines:
            dims = line["dimensions"]
            if dims["estimated"]:
                continue
            if (
                dims["length_m"] > profile["length_m"]
                or dims["width_m"] > profile["width_m"]
                or dims["height_m"] > profile["height_m"]
            ):
                oversize.append(line["name"])
        if oversize:
            continue
        manual_lines = [line for line in lines if manual_loads.get(line["id"], 0) > 0]
        automatic_lines = [line for line in lines if not manual_loads.get(line["id"], 0)]
        if manual_lines:
            manual_trips = max(
                math.ceil(line["quantity"] / manual_loads[line["id"]]) for line in manual_lines
            )
            automatic_weight = sum(line["weight_total_kg"] for line in automatic_lines)
            automatic_volume = sum(line["volume_total_m3"] for line in automatic_lines)
            weight_trips = math.ceil(automatic_weight / capacity_kg) if automatic_weight else 0
            volume_trips = (
                math.ceil(automatic_volume / profile["usable_volume_m3"])
                if automatic_volume else 0
            )
            trips = max(1, manual_trips, weight_trips, volume_trips)
            limiting_factor = "ручная загрузка"
        else:
            manual_trips = 0
            weight_trips = max(1, math.ceil(total_weight_kg / capacity_kg))
            volume_trips = max(1, math.ceil(total_volume_m3 / profile["usable_volume_m3"]))
            trips = max(weight_trips, volume_trips)
            limiting_factor = "вес" if weight_trips >= volume_trips else "габаритный объём"
        extra_km = 0 if is_moscow else max(0, distance - moscow_base_distance)
        trip_price = option["fixed_moscow_rub"] + extra_km * option["rate_rub_km"]
        alternatives.append({
            "vehicle": option["vehicle"],
            "capacity_t": option["capacity_t"],
            "rate_rub_km": option["rate_rub_km"],
            "fixed_moscow_rub": option["fixed_moscow_rub"],
            "moscow_base_distance": moscow_base_distance,
            "extra_km": extra_km,
            "pricing_kind": "фиксированный тариф по Москве" if is_moscow else (
                "фиксированный тариф + километры сверх пути до Москвы" if extra_km > 0
                else "фиксированный тариф (маршрут не длиннее пути до Москвы)"
            ),
            "trips": trips,
            "weight_trips": weight_trips,
            "volume_trips": volume_trips,
            "manual_trips": manual_trips,
            "manual_mode": bool(manual_lines),
            "delivery_total": trip_price * trips,
            "trip_price": trip_price,
            "loading_address": option["loading_address"],
            "profile": profile,
            "weight_load_pct": total_weight_kg / (capacity_kg * trips) * 100,
            "volume_load_pct": total_volume_m3 / (profile["usable_volume_m3"] * trips) * 100,
            "capacity_warning": total_weight_kg / trips > capacity_kg,
            "limiting_factor": limiting_factor,
        })
    if not alternatives:
        return None
    return min(alternatives, key=lambda item: (item["delivery_total"], item["trips"], -item["capacity_t"]))


def parse_cart(raw_value, product_by_id):
    try:
        raw_items = json.loads(raw_value or "[]")
    except (TypeError, ValueError, json.JSONDecodeError):
        raw_items = []
    merged = defaultdict(int)
    for item in raw_items if isinstance(raw_items, list) else []:
        product_id = str(item.get("id", ""))
        quantity = max(1, int(_float(item.get("quantity"), 1)))
        if product_id in product_by_id:
            merged[product_id] += quantity
    return [{"id": product_id, "quantity": quantity} for product_id, quantity in merged.items()]


def parse_vehicle_overrides(raw_value, delivery_rows):
    try:
        raw = json.loads(raw_value or "{}")
    except (TypeError, ValueError, json.JSONDecodeError):
        raw = {}
    if not isinstance(raw, dict):
        return {}
    allowed = {(row["supplier"], row["vehicle"]) for row in delivery_rows}
    return {
        str(supplier): str(vehicle)
        for supplier, vehicle in raw.items()
        if (str(supplier), str(vehicle)) in allowed
    }


def parse_manual_loads(raw_value, product_by_id):
    try:
        raw = json.loads(raw_value or "{}")
    except (TypeError, ValueError, json.JSONDecodeError):
        raw = {}
    if not isinstance(raw, dict):
        return {}
    result = {}
    for product_id, value in raw.items():
        product_id = str(product_id)
        quantity = int(_float(value, 0))
        if product_id in product_by_id and quantity > 0:
            result[product_id] = quantity
    return result


def build_quote(
    cart,
    product_by_id,
    delivery_rows,
    lat,
    lon,
    markup,
    is_moscow=False,
    vehicle_overrides=None,
    manual_loads=None,
):
    vehicle_overrides = vehicle_overrides or {}
    manual_loads = manual_loads or {}
    supplier_lines = defaultdict(list)
    all_lines = []
    for cart_item in cart:
        product = product_by_id[cart_item["id"]]
        quantity = cart_item["quantity"]
        dimensions = product["dimensions"]
        line = dict(product)
        line.update({
            "quantity": quantity,
            "dimensions": dimensions,
            "weight_total_kg": product["weight_kg"] * quantity,
            "volume_total_m3": dimensions["volume_m3"] * quantity,
            "purchase_total": product["price_rub"] * quantity,
            "delivery_share_total": 0,
            "delivery_unit": 0,
            "purchase_with_delivery_unit": product["price_rub"],
            "purchase_with_delivery_total": product["price_rub"] * quantity,
            "client_unit": product["price_rub"] * (1 + markup / 100),
            "client_line_total": product["price_rub"] * quantity * (1 + markup / 100),
            "manual_load_per_trip": manual_loads.get(product["id"], 0),
        })
        supplier_lines[product["supplier"]].append(line)
        all_lines.append(line)

    deliveries = []
    errors = []
    for supplier, lines in supplier_lines.items():
        options = [row for row in delivery_rows if row["supplier"] == supplier]
        if not options:
            errors.append(f"Для производителя «{supplier}» нет тарифа доставки.")
            continue
        distance, distance_kind = road_distance_km(options[0]["lat"], options[0]["lon"], lat, lon)
        moscow_base_distance, _ = road_distance_km(
            options[0]["lat"], options[0]["lon"], MOSCOW_CENTER_LAT, MOSCOW_CENTER_LON
        )
        selected_vehicle = vehicle_overrides.get(supplier)
        selected_options = [row for row in options if row["vehicle"] == selected_vehicle] if selected_vehicle else options
        calculation = mixed_delivery_calculation(
            lines, distance, selected_options, is_moscow, moscow_base_distance, manual_loads
        )
        if not calculation:
            errors.append(f"Заявка производителя «{supplier}» не помещается в доступный транспорт по весу или габаритам.")
            continue
        deliveries.append({
            "supplier": supplier,
            "lines": lines,
            "distance": distance,
            "distance_kind": distance_kind,
            "weight_total_kg": sum(line["weight_total_kg"] for line in lines),
            "volume_total_m3": sum(line["volume_total_m3"] for line in lines),
            **calculation,
        })

        supplier_weight = sum(line["weight_total_kg"] for line in lines)
        supplier_quantity = sum(line["quantity"] for line in lines)
        for line in lines:
            if supplier_weight > 0:
                delivery_ratio = line["weight_total_kg"] / supplier_weight
            else:
                delivery_ratio = line["quantity"] / supplier_quantity
            line["delivery_share_total"] = calculation["delivery_total"] * delivery_ratio
            line["delivery_unit"] = line["delivery_share_total"] / line["quantity"]
            line["purchase_with_delivery_unit"] = line["price_rub"] + line["delivery_unit"]
            line["purchase_with_delivery_total"] = line["purchase_with_delivery_unit"] * line["quantity"]
            line["client_unit"] = line["purchase_with_delivery_unit"] * (1 + markup / 100)
            line["client_line_total"] = line["client_unit"] * line["quantity"]
            if line["manual_load_per_trip"]:
                line["quantity_per_trip"] = f"до {line['manual_load_per_trip']} шт. (ручной режим)"
            else:
                minimum_per_trip = line["quantity"] // calculation["trips"]
                maximum_per_trip = math.ceil(line["quantity"] / calculation["trips"])
                if minimum_per_trip == maximum_per_trip:
                    line["quantity_per_trip"] = f"{maximum_per_trip} шт."
                elif minimum_per_trip == 0:
                    line["quantity_per_trip"] = f"до {maximum_per_trip} шт."
                else:
                    line["quantity_per_trip"] = f"{minimum_per_trip}–{maximum_per_trip} шт."

    purchase_total = sum(line["purchase_total"] for line in all_lines)
    delivery_total = sum(item["delivery_total"] for item in deliveries)
    purchase_with_delivery_total = sum(line["purchase_with_delivery_total"] for line in all_lines)
    client_total = sum(line["client_line_total"] for line in all_lines)
    return {
        "lines": all_lines,
        "deliveries": deliveries,
        "errors": errors,
        "purchase_total": purchase_total,
        "purchase_with_delivery_total": purchase_with_delivery_total,
        "delivery_total": delivery_total,
        "client_total": client_total,
        "weight_total_kg": sum(line["weight_total_kg"] for line in all_lines),
        "volume_total_m3": sum(line["volume_total_m3"] for line in all_lines),
    }


def build_full_load_catalog(products, delivery_options, lat, lon, markup, is_moscow=False):
    """Calculate one completely loaded vehicle for every product and vehicle type."""
    vehicles = []
    sorted_products = sorted(products, key=lambda item: (item["group"], item["name"], item["size_mm"]))
    for option in delivery_options:
        profile = vehicle_profile(option)
        capacity_kg = option["capacity_t"] * 1000
        distance, distance_kind = road_distance_km(option["lat"], option["lon"], lat, lon)
        moscow_base_distance, _ = road_distance_km(
            option["lat"], option["lon"], MOSCOW_CENTER_LAT, MOSCOW_CENTER_LON
        )
        extra_km = 0 if is_moscow else max(0, distance - moscow_base_distance)
        trip_price = option["fixed_moscow_rub"] + extra_km * option["rate_rub_km"]
        rows = []
        for product in sorted_products:
            dimensions = product["dimensions"]
            weight_kg = product["weight_kg"]
            price_rub = product["price_rub"]
            row = {
                **product,
                "load_quantity": 0,
                "load_weight_kg": 0,
                "purchase_goods_total": None,
                "purchase_with_delivery_unit": None,
                "client_unit": None,
                "client_total": None,
                "limiting_factor": "—",
                "status": "Расчёт выполнен",
            }
            if weight_kg <= 0:
                row["status"] = "Нет данных по массе"
                rows.append(row)
                continue
            if not dimensions["estimated"] and (
                dimensions["length_m"] > profile["length_m"]
                or dimensions["width_m"] > profile["width_m"]
                or dimensions["height_m"] > profile["height_m"]
            ):
                row["status"] = "Не помещается по габаритам"
                rows.append(row)
                continue
            by_weight = math.floor(capacity_kg / weight_kg)
            by_volume = math.floor(profile["usable_volume_m3"] / dimensions["volume_m3"])
            load_quantity = max(0, min(by_weight, by_volume))
            if load_quantity < 1:
                row["status"] = "Не помещается по весу или объёму"
                rows.append(row)
                continue
            row["load_quantity"] = load_quantity
            row["load_weight_kg"] = load_quantity * weight_kg
            row["limiting_factor"] = "масса" if by_weight <= by_volume else "габаритный объём"
            if dimensions["estimated"] and row["limiting_factor"] == "габаритный объём":
                row["limiting_factor"] = "оценочный объём"
            if price_rub <= 0:
                row["status"] = "Цена завода не указана"
                rows.append(row)
                continue
            row["purchase_goods_total"] = price_rub * load_quantity
            row["purchase_with_delivery_unit"] = price_rub + trip_price / load_quantity
            row["client_unit"] = row["purchase_with_delivery_unit"] * (1 + markup / 100)
            row["client_total"] = row["client_unit"] * load_quantity
            rows.append(row)
        vehicles.append({
            "vehicle": option["vehicle"],
            "capacity_t": option["capacity_t"],
            "profile": profile,
            "distance": distance,
            "distance_kind": distance_kind,
            "fixed_moscow_rub": option["fixed_moscow_rub"],
            "rate_rub_km": option["rate_rub_km"],
            "extra_km": extra_km,
            "trip_price": trip_price,
            "loading_address": option["loading_address"],
            "rows": rows,
            "calculated_count": sum(1 for row in rows if row["client_unit"] is not None),
            "unavailable_count": sum(1 for row in rows if row["client_unit"] is None),
        })
    return {
        "supplier": products[0]["supplier"] if products else "",
        "products_count": len(products),
        "vehicles": vehicles,
        "markup": markup,
    }


def _set_cell_fill(cell, color):
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _format_table(table, widths, header=True, font_size=9):
    """Apply deterministic full-width geometry and restrained business styling."""
    table.autofit = False
    table.alignment = 0
    width_dxa = [int(value * 1440) for value in widths]
    total_dxa = sum(width_dxa)
    table_properties = table._tbl.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(total_dxa))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_column, value in zip(grid.gridCol_lst, width_dxa):
        grid_column.set(qn("w:w"), str(value))
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.width = Inches(widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            cell_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            cell_width.set(qn("w:w"), str(width_dxa[column_index]))
            cell_width.set(qn("w:type"), "dxa")
            if header and row_index == 0:
                _set_cell_fill(cell, "E8EEF5")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 77, 120)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)


def _repeat_table_header(row):
    row_properties = row._tr.get_or_add_trPr()
    marker = row_properties.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        row_properties.append(marker)
    marker.set(qn("w:val"), "true")


def build_proposal_document(quote, form, markup):
    """Create a branded commercial proposal showing only client-facing prices."""
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(30, 91, 145)
        style.font.bold = True

    header = document.add_table(rows=1, cols=2)
    header.style = "Table Grid"
    logo_cell, company_cell = header.rows[0].cells
    logo_paragraph = logo_cell.paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO_B64_FILE.exists():
        logo_data = base64.b64decode(LOGO_B64_FILE.read_text(encoding="utf-8"))
        logo_paragraph.add_run().add_picture(BytesIO(logo_data), width=Cm(2.8))
    company_paragraph = company_cell.paragraphs[0]
    company_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    company_run = company_paragraph.add_run(COMPANY["short_name"])
    company_run.bold = True
    company_run.font.size = Pt(13)
    company_run.font.color.rgb = RGBColor(15, 86, 132)
    details = company_cell.add_paragraph(f'{COMPANY["phone"]} · {COMPANY["email"]}')
    details.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _format_table(header, [2.1, 4.95], header=False)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(15, 86, 132)
    subtitle = document.add_paragraph("Поставка железобетонных изделий с доставкой")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(92, 103, 117)

    client_rows = [
        ("Компания", form.get("client_company") or "—"),
        ("ИНН", form.get("client_inn") or "—"),
        ("Контактное лицо", form.get("client_contact_name") or "—"),
        ("Контакт", form.get("client_contact") or "—"),
        ("Адрес доставки", form.get("address") or "—"),
        ("Дата", date.today().strftime("%d.%m.%Y")),
    ]
    client_table = document.add_table(rows=0, cols=2)
    client_table.style = "Table Grid"
    for label, value in client_rows:
        cells = client_table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)
        cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_fill(cells[0], "F4F6F9")
    _format_table(client_table, [1.45, 5.6], header=False)

    heading = document.add_paragraph("Состав предложения", style="Heading 2")
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(5)
    products_table = document.add_table(rows=1, cols=6)
    products_table.style = "Table Grid"
    headers = ["№", "Наименование", "Производитель", "Кол-во", "Цена за ед. с доставкой", "Сумма"]
    for cell, value in zip(products_table.rows[0].cells, headers):
        cell.text = value
    for number, line in enumerate(quote["lines"], 1):
        cells = products_table.add_row().cells
        values = [
            number,
            f'{line["name"]}\n{line["size_mm"]}; {line["weight_kg"]:g} кг/шт.',
            line["supplier"],
            f'{line["quantity"]} шт.',
            f'{line["client_unit"]:,.0f} ₽'.replace(",", " "),
            f'{line["client_line_total"]:,.0f} ₽'.replace(",", " "),
        ]
        for cell, value in zip(cells, values):
            cell.text = str(value)
    _format_table(products_table, [0.35, 2.05, 1.25, 0.6, 1.35, 1.45])

    delivery_heading = document.add_paragraph("Доставка", style="Heading 2")
    delivery_heading.paragraph_format.space_before = Pt(10)
    delivery_heading.paragraph_format.space_after = Pt(5)
    delivery_table = document.add_table(rows=1, cols=5)
    delivery_table.style = "Table Grid"
    for cell, value in zip(delivery_table.rows[0].cells, ["Производитель", "Транспорт", "Рейсы", "Загрузка на рейс", "Стоимость"]):
        cell.text = value
    for item in quote["deliveries"]:
        cells = delivery_table.add_row().cells
        loading = "; ".join(f'{line["name"]}: {line["quantity_per_trip"]}' for line in item["lines"])
        values = [
            item["supplier"], item["vehicle"], item["trips"], loading,
            f'{item["delivery_total"]:,.0f} ₽'.replace(",", " "),
        ]
        for cell, value in zip(cells, values):
            cell.text = str(value)
    _format_table(delivery_table, [1.25, 1.1, 0.55, 2.7, 1.45])

    total = document.add_paragraph()
    total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total.paragraph_format.space_before = Pt(10)
    total_run = total.add_run(f'Итого для клиента: {quote["client_total"]:,.0f} ₽'.replace(",", " "))
    total_run.bold = True
    total_run.font.size = Pt(15)
    total_run.font.color.rgb = RGBColor(15, 86, 132)
    note = document.add_paragraph(
        f'Цена включает доставку до указанного объекта и наценку {markup:g}%. '
        "Окончательные сроки поставки и график рейсов согласовываются при подтверждении заказа."
    )
    note.paragraph_format.space_after = Pt(10)

    document.add_paragraph("Реквизиты поставщика", style="Heading 2")
    requisites = document.add_table(rows=0, cols=2)
    requisites.style = "Table Grid"
    for label, value in (
        ("Поставщик", COMPANY["full_name"]),
        ("ИНН / КПП", f'{COMPANY["inn"]} / {COMPANY["kpp"]}'),
        ("ОГРН", COMPANY["ogrn"]),
        ("Юридический адрес", COMPANY["address"]),
        ("Банк", COMPANY["bank"]),
        ("Р/с", COMPANY["account"]),
        ("К/с", COMPANY["correspondent_account"]),
        ("БИК", COMPANY["bik"]),
        ("Генеральный директор", COMPANY["director"]),
    ):
        cells = requisites.add_row().cells
        cells[0].text = label
        cells[1].text = value
        cells[0].paragraphs[0].runs[0].bold = True
    _format_table(requisites, [1.45, 5.6], header=False)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f'{COMPANY["short_name"]} · {COMPANY["phone"]} · {COMPANY["email"]}')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(100, 110, 122)
    return document


def build_catalog_proposal_document(offer, form):
    """Create a landscape proposal with every product and every supplier vehicle."""
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05
    for style_name, size in (("Heading 1", 16), ("Heading 2", 12), ("Heading 3", 10)):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(15, 86, 132)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True

    header = document.add_table(rows=1, cols=2)
    header.style = "Table Grid"
    logo_cell, company_cell = header.rows[0].cells
    logo_paragraph = logo_cell.paragraphs[0]
    if LOGO_B64_FILE.exists():
        logo_data = base64.b64decode(LOGO_B64_FILE.read_text(encoding="utf-8"))
        logo_paragraph.add_run().add_picture(BytesIO(logo_data), width=Cm(2.4))
    company_paragraph = company_cell.paragraphs[0]
    company_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    company_run = company_paragraph.add_run(COMPANY["short_name"])
    company_run.bold = True
    company_run.font.size = Pt(12)
    company_run.font.color.rgb = RGBColor(15, 86, 132)
    contact = company_cell.add_paragraph(f'{COMPANY["phone"]} · {COMPANY["email"]}')
    contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _format_table(header, [2.0, 7.9], header=False, font_size=8.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(1)
    title_run = title.add_run("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(15, 86, 132)
    subtitle = document.add_paragraph(
        f'Полный каталог поставок {offer["supplier"]}: варианты полной загрузки машин'
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(7)
    subtitle.runs[0].font.size = Pt(10)
    subtitle.runs[0].font.color.rgb = RGBColor(92, 103, 117)

    metadata = document.add_table(rows=0, cols=4)
    metadata.style = "Table Grid"
    metadata_rows = [
        ("Компания клиента", form.get("client_company") or "—", "ИНН", form.get("client_inn") or "—"),
        ("Контактное лицо", form.get("client_contact_name") or "—", "Контакт", form.get("client_contact") or "—"),
        ("Адрес доставки", form.get("address") or "—", "Дата", date.today().strftime("%d.%m.%Y")),
        ("Производитель", offer["supplier"], "Наценка менеджера", f'{offer["markup"]:g}%'),
    ]
    for left_label, left_value, right_label, right_value in metadata_rows:
        cells = metadata.add_row().cells
        for cell, value in zip(cells, (left_label, left_value, right_label, right_value)):
            cell.text = str(value)
        for index in (0, 2):
            cells[index].paragraphs[0].runs[0].bold = True
            _set_cell_fill(cells[index], "F1F5F9")
    _format_table(metadata, [1.2, 4.05, 1.2, 3.45], header=False, font_size=8)

    intro = document.add_paragraph(
        "В таблицах ниже каждая строка — отдельный вариант поставки одной полной машиной. "
        "Цена клиента включает стоимость изделий, доставку до указанного адреса и наценку менеджера."
    )
    intro.paragraph_format.space_before = Pt(6)
    intro.paragraph_format.space_after = Pt(4)

    for vehicle_index, vehicle in enumerate(offer["vehicles"]):
        if vehicle_index:
            document.add_page_break()
        heading = document.add_paragraph(
            f'{vehicle["vehicle"]} · грузоподъёмность {vehicle["capacity_t"]:g} т',
            style="Heading 2",
        )
        heading.paragraph_format.space_after = Pt(2)
        tariff = document.add_paragraph(
            f'Доставка одной машины: {vehicle["trip_price"]:,.0f} ₽ · '
            f'расстояние {vehicle["distance"]:.1f} км ({vehicle["distance_kind"]}) · '
            f'рассчитано {vehicle["calculated_count"]} из {offer["products_count"]} позиций. '
            f'Погрузка: {vehicle["loading_address"]}'.replace(",", " ")
        )
        tariff.paragraph_format.space_after = Pt(4)
        tariff.runs[0].font.size = Pt(8.5)

        table = document.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        headers = [
            "№", "Изделие / раздел", "Габариты / масса", "Полная загрузка",
            "Цена завода", "Доставка / рейс", "Клиенту за 1 ед.", "Полная машина / статус",
        ]
        for cell, value in zip(table.rows[0].cells, headers):
            cell.text = value
        _repeat_table_header(table.rows[0])
        for number, row in enumerate(vehicle["rows"], 1):
            cells = table.add_row().cells
            dimensions = row["dimensions"]
            dimension_note = f'{row["size_mm"]}; {row["weight_kg"]:g} кг'
            if dimensions["estimated"]:
                dimension_note += " (объём оценочный)"
            if row["client_unit"] is not None:
                final_value = (
                    f'{row["client_total"]:,.0f} ₽\n'
                    f'Загрузка: {row["load_weight_kg"] / 1000:.2f} т; лимит: {row["limiting_factor"]}'
                )
                values = [
                    number,
                    f'{row["name"]}\n{row["group"]}',
                    dimension_note,
                    f'{row["load_quantity"]} шт.',
                    f'{row["price_rub"]:,.0f} ₽/шт.',
                    f'{vehicle["trip_price"]:,.0f} ₽',
                    f'{row["client_unit"]:,.0f} ₽',
                    final_value,
                ]
            else:
                values = [
                    number, f'{row["name"]}\n{row["group"]}', dimension_note, "—",
                    f'{row["price_rub"]:,.0f} ₽/шт.' if row["price_rub"] > 0 else "—",
                    f'{vehicle["trip_price"]:,.0f} ₽', "—", row["status"],
                ]
            for cell, value in zip(cells, values):
                cell.text = str(value).replace(",", " ")
        _format_table(table, [0.32, 2.35, 1.22, 0.78, 0.9, 0.92, 1.1, 2.31], font_size=7.2)

    document.add_paragraph("Условия предложения", style="Heading 2")
    document.add_paragraph(
        "Расчёт носит коммерческий характер. Фактическую схему укладки, допустимую загрузку по осям, "
        "наличие изделий и сроки отгрузки необходимо подтвердить у производителя перед заказом. "
        "Для позиций с оценочным объёмом требуется дополнительная проверка габаритов."
    )
    document.add_paragraph("Реквизиты поставщика", style="Heading 2")
    requisites = document.add_table(rows=0, cols=4)
    requisites.style = "Table Grid"
    for values in (
        ("Поставщик", COMPANY["full_name"], "ИНН / КПП", f'{COMPANY["inn"]} / {COMPANY["kpp"]}'),
        ("ОГРН", COMPANY["ogrn"], "Банк", COMPANY["bank"]),
        ("Р/с", COMPANY["account"], "БИК", COMPANY["bik"]),
        ("Генеральный директор", COMPANY["director"], "Юридический адрес", COMPANY["address"]),
    ):
        cells = requisites.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = str(value)
        for index in (0, 2):
            cells[index].paragraphs[0].runs[0].bold = True
    _format_table(requisites, [1.3, 3.65, 1.3, 3.65], header=False, font_size=7.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f'{COMPANY["short_name"]} · {COMPANY["phone"]} · {COMPANY["email"]}')
    footer_run.font.size = Pt(7.5)
    footer_run.font.color.rgb = RGBColor(100, 110, 122)
    return document


CATALOG_PAGE = r"""
<!doctype html><html lang="ru"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>КП полного каталога ЖБИ</title><style>
:root{--ink:#18202a;--muted:#647184;--blue:#185bd8;--line:#dce3ec;--bg:#f3f6fa;--ok:#e9f8ef;--warn:#fff6df}
*{box-sizing:border-box}body{font-family:Arial,sans-serif;background:var(--bg);color:var(--ink);margin:0;padding:24px}.wrap{max-width:1180px;margin:auto}.card{background:#fff;border-radius:16px;padding:22px;margin-bottom:18px;box-shadow:0 5px 18px #20305012}.nav a{margin-right:18px;color:var(--blue);font-weight:700;text-decoration:none}h1,h2,h3{margin-top:0}.muted{color:var(--muted);font-size:13px}.grid{display:grid;grid-template-columns:repeat(4,minmax(0,1fr));gap:14px}.span2{grid-column:span 2}label{display:block;font-size:13px;font-weight:700;margin-bottom:6px}input,select,button{width:100%;min-height:44px;border:1px solid var(--line);border-radius:10px;padding:10px 12px;font-size:15px;background:#fff}button{background:var(--blue);border-color:var(--blue);color:#fff;font-weight:700;cursor:pointer}.secondary{background:#fff;color:var(--blue)}.summary{background:var(--ok);border:1px solid #bfe5cc}.warning{background:var(--warn);border:1px solid #f0d58a}.vehicle{border:1px solid var(--line);border-radius:12px;padding:15px;margin-top:12px}.badges{display:flex;gap:8px;flex-wrap:wrap}.badge{background:#eef3ff;border-radius:999px;padding:6px 10px;font-size:12px;font-weight:700}.actions{display:flex;gap:12px;justify-content:flex-end;margin-top:18px}.actions button{max-width:420px}@media(max-width:800px){body{padding:12px}.grid{grid-template-columns:1fr}.span2{grid-column:span 1}.actions{display:block}.actions button{max-width:none;margin-top:10px}}
</style></head><body><div class="wrap">
<div class="card nav"><a href="/">Нерудные материалы</a><a href="/zbi">Калькулятор ЖБИ</a><a href="/zbi/catalog-proposal">КП полного каталога</a><a href="/carriers">Перевозчики</a></div>
<form method="post"><div class="card"><h1>КП по полной загрузке каталога производителя</h1><p class="muted">Выберите одного производителя. Система рассчитает поставку каждого его изделия отдельной полной машиной для всех доступных видов транспорта и подготовит единый Word‑документ.</p>
<div class="grid"><div class="span2"><label>Производитель</label><select name="supplier" required><option value="">Выберите производителя</option>{% for item in suppliers %}<option value="{{ item }}"{% if form.supplier == item %} selected{% endif %}>{{ item }}</option>{% endfor %}</select></div><div class="span2"><label>Адрес доставки</label><input name="address" value="{{ form.address }}" placeholder="Москва, улица и дом" required></div><div><label>Наценка менеджера, %</label><input name="markup" type="number" min="0" step="0.1" value="{{ form.markup }}" required></div><div><label>Компания клиента</label><input name="client_company" value="{{ form.client_company }}"></div><div><label>ИНН клиента</label><input name="client_inn" value="{{ form.client_inn }}"></div><div><label>ФИО контактного лица</label><input name="client_contact_name" value="{{ form.client_contact_name }}"></div><div class="span2"><label>Телефон или e-mail</label><input name="client_contact" value="{{ form.client_contact }}"></div></div>
<div class="actions"><button type="submit">Рассчитать полный каталог</button>{% if offer %}<button type="submit" class="secondary" formaction="/zbi/catalog-proposal.docx" formmethod="post">Скачать КП в Word</button>{% endif %}</div></div></form>
{% if error %}<div class="card warning"><b>Расчёт не выполнен.</b> {{ error }}</div>{% endif %}
{% if offer %}<div class="card summary"><h2>{{ offer.supplier }}</h2><p><b>{{ offer.products_count }} изделий</b> · наценка {{ offer.markup }}% · рассчитано для {{ offer.vehicles|length }} видов транспорта.</p>{% for vehicle in offer.vehicles %}<div class="vehicle"><h3>{{ vehicle.vehicle }} · {{ vehicle.capacity_t|round(0)|int }} т</h3><div class="badges"><span class="badge">рейс {{ vehicle.trip_price|money }} ₽</span><span class="badge">{{ vehicle.distance|round(1) }} км</span><span class="badge">рассчитано {{ vehicle.calculated_count }}</span><span class="badge">нужна проверка {{ vehicle.unavailable_count }}</span></div><p class="muted">В Word‑документ войдут все {{ offer.products_count }} позиций, включая строки с отсутствующими данными или превышением габаритов.</p></div>{% endfor %}</div>{% endif %}
</div></body></html>
"""


PAGE = r"""
<!doctype html><html lang="ru"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Калькулятор заявки ЖБИ</title>
<style>
:root{--ink:#18202a;--muted:#647184;--blue:#185bd8;--line:#dce3ec;--bg:#f3f6fa;--ok:#e9f8ef;--danger:#a62d2d}
*{box-sizing:border-box}body{font-family:Arial,sans-serif;background:var(--bg);color:var(--ink);margin:0;padding:24px}.wrap{max-width:1380px;margin:auto}.card{background:#fff;border-radius:16px;padding:22px;margin-bottom:18px;box-shadow:0 5px 18px #20305012}h1,h2,h3{margin-top:0}.nav a{margin-right:18px;color:var(--blue);font-weight:700;text-decoration:none}.grid{display:grid;grid-template-columns:repeat(4,minmax(0,1fr));gap:14px}.span2{grid-column:span 2}.span4{grid-column:span 4}label{display:block;font-size:13px;font-weight:700;margin-bottom:6px}input,select,button{width:100%;min-height:44px;border:1px solid var(--line);border-radius:10px;padding:10px 12px;font-size:15px;background:#fff}button{background:var(--blue);border-color:var(--blue);color:#fff;font-weight:700;cursor:pointer}.secondary{background:#fff;color:var(--blue)}.remove{background:#fff;color:var(--danger);border-color:#e6bcbc;padding:7px;min-height:34px}.hint,.muted{color:var(--muted);font-size:13px}.suggestions{max-height:360px;overflow:auto;background:#fff;border:1px solid var(--line);border-radius:10px}.suggestion{display:block;width:100%;min-height:0;padding:10px 12px;background:#fff;color:var(--ink);text-align:left;border:0;border-bottom:1px solid #edf0f4;border-radius:0;cursor:pointer}.suggestion:hover,.suggestion.selected{background:#eef4ff}.suggestion small{display:block;color:var(--muted);margin-top:3px;font-weight:400}.list-head{display:flex;justify-content:space-between;gap:12px;align-items:center;margin-bottom:7px}.transport-grid{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:12px;margin-top:14px}.transport-choice{padding:12px;border:1px solid var(--line);border-radius:12px;background:#f8faff}.calculate-footer{display:flex;justify-content:flex-end;gap:12px;flex-wrap:wrap;margin-top:20px;padding-top:18px;border-top:1px solid var(--line)}.calculate-footer button{max-width:420px}.summary{background:var(--ok);border:1px solid #bfe5cc}.warning{background:#fff6df;border:1px solid #f0d58a}.kpis{display:grid;grid-template-columns:repeat(4,1fr);gap:12px}.kpi{background:#f5f8fc;border-radius:12px;padding:14px}.kpi b{display:block;font-size:22px;margin-top:4px}.delivery{border:1px solid var(--line);border-radius:14px;padding:16px;margin-top:12px}.badges{display:flex;flex-wrap:wrap;gap:7px;margin:8px 0}.badge{background:#eef3ff;border-radius:999px;padding:6px 9px;font-size:12px;font-weight:700}.table-wrap{overflow:auto}table{width:100%;border-collapse:collapse;margin-top:12px}th,td{text-align:left;padding:10px;border-bottom:1px solid #e6ebf1;vertical-align:top}th{font-size:12px;color:var(--muted)}.money{font-weight:800;white-space:nowrap}.empty{text-align:center;color:var(--muted);padding:24px}.actions{display:flex;gap:10px;align-items:end}.actions>*{flex:1}
@media(max-width:900px){body{padding:12px}.grid{grid-template-columns:1fr}.span2,.span4{grid-column:span 1}.kpis{grid-template-columns:1fr 1fr}.actions{display:block}.actions>*{margin-top:8px}}
</style></head><body><div class="wrap">
<div class="card nav"><a href="/">Нерудные материалы</a><a href="/zbi">Калькулятор ЖБИ</a><a href="/zbi/catalog-proposal">КП полного каталога</a><a href="/carriers">Перевозчики</a></div>
<form method="post" id="quoteForm"><input type="hidden" name="items_json" id="itemsJson"><input type="hidden" name="vehicles_json" id="vehiclesJson"><input type="hidden" name="loads_json" id="loadsJson">
<div class="card"><h1>Расчёт и формирование заявки ЖБИ</h1><p class="muted">Добавьте изделия, укажите адрес и наценку. Для каждой позиции будет рассчитана закупочная цена с доставкой и цена клиенту.</p>
<div class="grid"><div class="span2"><label>Адрес доставки</label><input name="address" value="{{ form.address }}" placeholder="Москва, улица и дом" required></div><div class="span2"><label>Наценка на полную стоимость, %</label><input type="number" min="0" step="0.1" name="markup" value="{{ form.markup }}" required></div></div>
<h2 style="margin-top:24px">Добавить изделие</h2><div class="grid">
<div><label>Производитель</label><select id="supplier"><option value="">Любой производитель</option>{% for item in suppliers %}<option value="{{ item }}">{{ item }}</option>{% endfor %}</select></div>
<div><label>Раздел</label><select id="group"><option value="">Все изделия / любые</option>{% for item in groups %}<option value="{{ item }}">{{ item }} — любые</option>{% endfor %}</select></div>
<div class="span2"><label>Поиск изделия</label><input id="productSearch" autocomplete="off" placeholder="Например: ФБС 24.4.6, 2П 30.18 или лоток"></div>
<div><label>Количество, шт.</label><input id="addQuantity" type="number" min="1" step="1" value="1"></div><div class="span2"><label>Выбрано</label><input id="selectedName" readonly placeholder="Сначала выберите позицию из списка ниже"></div><div><label>&nbsp;</label><button type="button" id="addItem">Добавить в заявку</button></div>
<div class="span4"><div class="list-head"><label style="margin:0">Выбор изделия из раздела</label><span id="productCount" class="muted"></span></div><div id="suggestions" class="suggestions"></div></div>
</div>
<h2 style="margin-top:24px">Состав заявки</h2><div class="table-wrap"><table><thead><tr><th>Изделие</th><th>Производитель</th><th>Габариты</th><th>Масса 1 шт.</th><th>Количество</th><th>Загрузка на 1 машину</th><th>Общая масса</th><th>Цена завода без доставки</th><th></th></tr></thead><tbody id="cartBody"></tbody></table></div><div id="emptyCart" class="empty">В заявке пока нет изделий</div><div id="transportChoices" class="transport-grid"></div><p id="transportHint" class="muted">Транспорт появится после добавления изделия. Поле «Загрузка на 1 машину» оставьте пустым для автоматического расчёта или задайте своё количество.</p><p class="muted"><b>Расчёт не запускается автоматически.</b> Сначала спокойно добавьте все позиции и настройте загрузку, затем нажмите кнопку ниже.</p>{% if quote %}<h2 style="margin-top:24px">Данные клиента для коммерческого предложения</h2><div class="grid"><div><label>Наименование компании</label><input name="client_company" value="{{ form.client_company }}"></div><div><label>ИНН</label><input name="client_inn" value="{{ form.client_inn }}"></div><div><label>ФИО контактного лица</label><input name="client_contact_name" value="{{ form.client_contact_name }}"></div><div><label>Контакт</label><input name="client_contact" value="{{ form.client_contact }}" placeholder="Телефон или e-mail"></div></div>{% endif %}<p id="pendingChanges" class="warning" style="display:none;padding:10px;border-radius:10px;margin-top:16px"><b>Заявка изменена.</b> Нажмите «Пересчитать заявку», чтобы обновить цены и количество рейсов.</p><div class="calculate-footer"><button type="submit">{{ 'Пересчитать заявку' if quote else 'Начать расчёт заявки' }}</button>{% if quote %}<button type="submit" class="secondary" formaction="/zbi/proposal.docx" formmethod="post">Скачать коммерческое предложение</button>{% endif %}</div></div>
</form>
{% if error %}<div class="card warning"><b>Не удалось выполнить расчёт.</b> {{ error }}</div>{% endif %}
{% if quote %}
<div class="card summary"><h2>Итог заявки</h2><div class="kpis"><div class="kpi">Товары по закупке<b>{{ quote.purchase_total|money }} ₽</b><span class="muted">без доставки</span></div><div class="kpi">Закупка с доставкой<b>{{ quote.purchase_with_delivery_total|money }} ₽</b><span class="muted">товары + доставка</span></div><div class="kpi">Доставка отдельно<b>{{ quote.delivery_total|money }} ₽</b><span class="muted">до объекта</span></div><div class="kpi">Цена клиенту<b>{{ quote.client_total|money }} ₽</b><span class="muted">с наценкой {{ markup }}% на полную стоимость</span></div></div><p><b>Общая масса:</b> {{ quote.weight_total_kg|round(0)|int }} кг · <b>транспортный габаритный объём:</b> {{ quote.volume_total_m3|round(2) }} м³</p></div>
<div class="card"><h2>Стоимость каждой позиции</h2><p class="muted">Доставка распределяется между изделиями одного производителя пропорционально их массе.</p><div class="table-wrap"><table><thead><tr><th>Изделие</th><th>Кол-во</th><th>Габариты и масса</th><th>Завод за 1 шт.</th><th>Доставка на 1 шт.</th><th>Закупка за 1 шт. с доставкой</th><th>Цена клиенту за 1 шт.</th><th>Клиенту всего</th></tr></thead><tbody>{% for line in quote.lines %}<tr><td><b>{{ line.name }}</b><div class="muted">{{ line.supplier }} · {{ line.group }}</div></td><td>{{ line.quantity }}</td><td>{{ line.size_mm }}<br>{{ line.weight_kg|round(1) }} кг/шт.{% if line.dimensions.estimated %}<div class="muted">объём оценён по массе</div>{% endif %}</td><td class="money">{{ line.price_rub|money }} ₽</td><td class="money">{{ line.delivery_unit|money }} ₽</td><td class="money">{{ line.purchase_with_delivery_unit|money }} ₽</td><td class="money">{{ line.client_unit|money }} ₽<div class="muted">наценка {{ markup }}%</div></td><td class="money">{{ line.client_line_total|money }} ₽</td></tr>{% endfor %}</tbody></table></div></div>
<div class="card"><h2>Доставка до объекта — отдельно</h2>{% for item in quote.deliveries %}<div class="delivery"><h3>{{ item.supplier }}</h3><div class="badges"><span class="badge">{{ item.vehicle }} · {{ item.capacity_t|round(0)|int }} т</span><span class="badge">{{ item.trips }} рейс(а)</span><span class="badge">{{ item.distance|round(1) }} км · {{ item.distance_kind }}</span><span class="badge">{{ item.pricing_kind }}</span><span class="badge">ограничение: {{ item.limiting_factor }}</span></div>{% if item.capacity_warning %}<div class="warning" style="padding:10px;border-radius:10px"><b>Проверьте ручную загрузку:</b> заданное количество даёт превышение грузоподъёмности выбранной машины.</div>{% endif %}<p><b>{{ item.delivery_total|money }} ₽ за доставку</b> · фиксированно до Москвы {{ item.fixed_moscow_rub|money }} ₽/рейс{% if item.extra_km > 0 %} + {{ item.extra_km|round(1) }} км × {{ item.rate_rub_km|money }} ₽/км{% endif %} · масса {{ item.weight_total_kg|round(0)|int }} кг · габаритный объём {{ item.volume_total_m3|round(2) }} м³</p><div class="table-wrap"><table><thead><tr><th>Изделие</th><th>Всего</th><th>На один рейс</th></tr></thead><tbody>{% for line in item.lines %}<tr><td>{{ line.name }}</td><td>{{ line.quantity }} шт.</td><td><b>{{ line.quantity_per_trip }}</b></td></tr>{% endfor %}</tbody></table></div><p class="muted">Средняя загрузка одного рейса: по массе {{ item.weight_load_pct|round(0)|int }}%, по объёму {{ item.volume_load_pct|round(0)|int }}%. Погрузка: {{ item.loading_address }}</p></div>{% endfor %}{% for message in quote.errors %}<div class="warning">{{ message }}</div>{% endfor %}</div>
{% endif %}
</div><script>
const catalog={{ catalog_json|safe }},deliveryCatalog={{ delivery_json|safe }};
let cart={{ cart_json|safe }},vehicleOverrides={{ vehicle_overrides_json|safe }},manualLoads={{ manual_loads_json|safe }},selectedId='';
const hasQuote={{ 'true' if quote else 'false' }};
const supplier=document.getElementById('supplier'),group=document.getElementById('group'),search=document.getElementById('productSearch'),box=document.getElementById('suggestions'),productCount=document.getElementById('productCount'),selectedName=document.getElementById('selectedName'),qty=document.getElementById('addQuantity'),body=document.getElementById('cartBody'),empty=document.getElementById('emptyCart'),itemsJson=document.getElementById('itemsJson'),vehiclesJson=document.getElementById('vehiclesJson'),loadsJson=document.getElementById('loadsJson'),transportChoices=document.getElementById('transportChoices'),transportHint=document.getElementById('transportHint'),pendingChanges=document.getElementById('pendingChanges'),quoteForm=document.getElementById('quoteForm');
const esc=s=>String(s).replace(/[&<>'"]/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;',"'":'&#39;','"':'&quot;'}[c]));
function sync(){itemsJson.value=JSON.stringify(cart);vehiclesJson.value=JSON.stringify(vehicleOverrides);loadsJson.value=JSON.stringify(manualLoads)}
function scheduleRecalculate(){sync();if(pendingChanges&&hasQuote)pendingChanges.style.display='block'}
function filtered(){const q=search.value.trim().toLowerCase();if(!group.value&&q.length<2)return[];return catalog.filter(p=>(!supplier.value||p.supplier===supplier.value)&&(!group.value||p.group===group.value)&&(!q||p.name.toLowerCase().includes(q)||p.size_mm.toLowerCase().includes(q))).slice(0,500)}
function show(){const rows=filtered();productCount.textContent=rows.length?`${rows.length} позиций`:(group.value||search.value.trim().length>=2?'0 позиций':'Выберите раздел или введите минимум 2 символа');box.innerHTML=rows.map(p=>`<button type="button" class="suggestion${p.id===selectedId?' selected':''}" data-id="${p.id}"><b>${esc(p.name)}</b><small>${esc(p.supplier)} · ${esc(p.group)} · ${esc(p.size_mm)} · ${p.weight_kg||'масса не указана'} кг · ${p.price_rub} ₽</small></button>`).join('')||'<div class="empty">Изделия появятся здесь после выбора раздела</div>';box.querySelectorAll('[data-id]').forEach(el=>el.onclick=()=>{const p=catalog.find(x=>x.id===el.dataset.id);selectedId=p.id;selectedName.value=`${p.name} — ${p.supplier}`;show()})}
function resetSelection(clearSearch=true){selectedId='';selectedName.value='';if(clearSearch)search.value='';show()}
supplier.addEventListener('change',()=>resetSelection());group.addEventListener('change',()=>resetSelection());search.addEventListener('input',()=>resetSelection(false));
function renderTransport(){const suppliers=[...new Set(cart.map(item=>catalog.find(p=>p.id===item.id)?.supplier).filter(Boolean))];Object.keys(vehicleOverrides).forEach(name=>{if(!suppliers.includes(name))delete vehicleOverrides[name]});transportChoices.innerHTML=suppliers.map(name=>{const options=deliveryCatalog.filter(row=>row.supplier===name);if(!options.length)return`<div class="transport-choice"><b>${esc(name)}</b><div class="muted">Нет тарифа доставки</div></div>`;if(!options.some(row=>row.vehicle===vehicleOverrides[name]))vehicleOverrides[name]=options[0].vehicle;return`<div class="transport-choice"><label>Транспорт: ${esc(name)}</label><select class="vehicleChoice" data-supplier="${esc(name)}">${options.map(row=>`<option value="${esc(row.vehicle)}"${row.vehicle===vehicleOverrides[name]?' selected':''}>${esc(row.vehicle)} · ${row.capacity_t} т · Москва ${Math.round(row.fixed_moscow_rub).toLocaleString('ru-RU')} ₽ · далее ${row.rate_rub_km} ₽/км</option>`).join('')}</select></div>`}).join('');transportHint.style.display=suppliers.length?'none':'block';transportChoices.querySelectorAll('.vehicleChoice').forEach(el=>el.onchange=()=>{vehicleOverrides[el.dataset.supplier]=el.value;sync();scheduleRecalculate()});sync()}
function render(){body.innerHTML=cart.map((item,i)=>{const p=catalog.find(x=>x.id===item.id);const total=p.weight_kg*item.quantity;return`<tr><td><b>${esc(p.name)}</b><div class="muted">${esc(p.group)}</div></td><td>${esc(p.supplier)}</td><td>${esc(p.size_mm)}</td><td>${p.weight_kg} кг</td><td><input class="cartQty" data-index="${i}" type="number" min="1" step="1" value="${item.quantity}"></td><td><input class="manualLoad" data-id="${p.id}" type="number" min="1" step="1" value="${manualLoads[p.id]||''}" placeholder="Авто"></td><td><b class="lineWeight" data-index="${i}">${Math.round(total)} кг</b></td><td class="money">${Math.round(p.price_rub).toLocaleString('ru-RU')} ₽/шт.</td><td><button type="button" class="remove" data-index="${i}">Удалить</button></td></tr>`}).join('');empty.style.display=cart.length?'none':'block';renderTransport();document.querySelectorAll('.cartQty').forEach(el=>{el.oninput=()=>{const index=+el.dataset.index;cart[index].quantity=Math.max(1,parseInt(el.value)||1);const p=catalog.find(x=>x.id===cart[index].id),weight=document.querySelector(`.lineWeight[data-index="${index}"]`);if(weight)weight.textContent=`${Math.round(p.weight_kg*cart[index].quantity)} кг`;sync();scheduleRecalculate()};el.onblur=()=>{el.value=cart[+el.dataset.index].quantity}});document.querySelectorAll('.manualLoad').forEach(el=>el.oninput=()=>{const value=parseInt(el.value)||0;if(value>0)manualLoads[el.dataset.id]=value;else delete manualLoads[el.dataset.id];sync();scheduleRecalculate()});document.querySelectorAll('.remove').forEach(el=>el.onclick=()=>{const removed=cart[+el.dataset.index];if(removed)delete manualLoads[removed.id];cart.splice(+el.dataset.index,1);render();scheduleRecalculate()})}
document.getElementById('addItem').onclick=()=>{if(!selectedId){alert('Выберите точное изделие из списка ниже');return}const amount=Math.max(1,parseInt(qty.value)||1);const old=cart.find(x=>x.id===selectedId);if(old)old.quantity+=amount;else cart.push({id:selectedId,quantity:amount});qty.value=1;resetSelection();render();scheduleRecalculate()};
quoteForm.elements.address.addEventListener('change',scheduleRecalculate);quoteForm.elements.markup.addEventListener('change',scheduleRecalculate);quoteForm.onsubmit=e=>{if(!cart.length){e.preventDefault();alert('Добавьте хотя бы одно изделие в заявку');return}sync()};show();render();
</script></body></html>
"""


@zbi_bp.app_template_filter("money")
def money(value):
    return f"{float(value):,.0f}".replace(",", " ")


@zbi_bp.route("", methods=["GET", "POST"])
@zbi_bp.route("/", methods=["GET", "POST"])
def calculator():
    products, delivery_rows = load_catalog()
    product_by_id = {item["id"]: item for item in products}
    form = {
        "address": request.form.get("address", ""),
        "markup": request.form.get("markup", "10"),
        "client_company": request.form.get("client_company", ""),
        "client_inn": request.form.get("client_inn", ""),
        "client_contact_name": request.form.get("client_contact_name", ""),
        "client_contact": request.form.get("client_contact", ""),
    }
    markup = max(0, _float(form["markup"], 0))
    cart = parse_cart(request.form.get("items_json", "[]"), product_by_id)
    vehicle_overrides = parse_vehicle_overrides(request.form.get("vehicles_json", "{}"), delivery_rows)
    manual_loads = parse_manual_loads(request.form.get("loads_json", "{}"), product_by_id)
    quote = None
    error = ""
    if request.method == "POST":
        if not cart:
            error = "Добавьте хотя бы одно изделие в заявку."
        else:
            lat, lon, is_moscow = geocode(form["address"])
            if lat is None:
                error = "Адрес не найден. Проверьте адрес или введите координаты через запятую."
            else:
                quote = build_quote(
                    cart, product_by_id, delivery_rows, lat, lon, markup, is_moscow,
                    vehicle_overrides, manual_loads
                )
                if not quote["deliveries"]:
                    error = "Не удалось подобрать транспорт ни для одного производителя."
    catalog_for_js = [{
        key: item[key] for key in ("id", "supplier", "group", "name", "size_mm", "weight_kg", "price_rub")
    } for item in products]
    delivery_for_js = [{
        key: item[key] for key in ("supplier", "vehicle", "capacity_t", "fixed_moscow_rub", "rate_rub_km")
    } for item in delivery_rows]
    return render_template_string(
        PAGE,
        suppliers=sorted({item["supplier"] for item in products}),
        groups=sorted({item["group"] for item in products}),
        form=form,
        markup=markup,
        quote=quote,
        error=error,
        catalog_json=json.dumps(catalog_for_js, ensure_ascii=False).replace("</", "<\\/"),
        cart_json=json.dumps(cart, ensure_ascii=False).replace("</", "<\\/"),
        delivery_json=json.dumps(delivery_for_js, ensure_ascii=False).replace("</", "<\\/"),
        vehicle_overrides_json=json.dumps(vehicle_overrides, ensure_ascii=False).replace("</", "<\\/"),
        manual_loads_json=json.dumps(manual_loads, ensure_ascii=False).replace("</", "<\\/"),
    )


def _catalog_proposal_input():
    products, delivery_rows = load_catalog()
    suppliers_with_delivery = {item["supplier"] for item in delivery_rows}
    suppliers = sorted({item["supplier"] for item in products} & suppliers_with_delivery)
    form = {
        "supplier": request.form.get("supplier", "").strip(),
        "address": request.form.get("address", "").strip(),
        "markup": request.form.get("markup", "10"),
        "client_company": request.form.get("client_company", "").strip(),
        "client_inn": request.form.get("client_inn", "").strip(),
        "client_contact_name": request.form.get("client_contact_name", "").strip(),
        "client_contact": request.form.get("client_contact", "").strip(),
    }
    markup = max(0, _float(form["markup"], 0))
    if form["supplier"] not in suppliers:
        return suppliers, form, markup, None, "Выберите производителя из списка."
    lat, lon, is_moscow = geocode(form["address"])
    if lat is None:
        return suppliers, form, markup, None, "Адрес не найден. Проверьте улицу и номер дома."
    supplier_products = [item for item in products if item["supplier"] == form["supplier"]]
    supplier_delivery = [item for item in delivery_rows if item["supplier"] == form["supplier"]]
    offer = build_full_load_catalog(
        supplier_products, supplier_delivery, lat, lon, markup, is_moscow
    )
    if not offer["vehicles"]:
        return suppliers, form, markup, None, "Для производителя не найдены тарифы транспорта."
    return suppliers, form, markup, offer, ""


@zbi_bp.route("/catalog-proposal", methods=["GET", "POST"])
def catalog_proposal():
    products, delivery_rows = load_catalog()
    suppliers = sorted(
        {item["supplier"] for item in products}
        & {item["supplier"] for item in delivery_rows}
    )
    form = {
        "supplier": request.form.get("supplier", "").strip(),
        "address": request.form.get("address", "").strip(),
        "markup": request.form.get("markup", "10"),
        "client_company": request.form.get("client_company", "").strip(),
        "client_inn": request.form.get("client_inn", "").strip(),
        "client_contact_name": request.form.get("client_contact_name", "").strip(),
        "client_contact": request.form.get("client_contact", "").strip(),
    }
    offer = None
    error = ""
    markup = max(0, _float(form["markup"], 0))
    if request.method == "POST":
        suppliers, form, markup, offer, error = _catalog_proposal_input()
    return render_template_string(
        CATALOG_PAGE,
        suppliers=suppliers,
        form=form,
        markup=markup,
        offer=offer,
        error=error,
    )


@zbi_bp.route("/catalog-proposal.docx", methods=["POST"])
def catalog_proposal_docx():
    _, form, _, offer, error = _catalog_proposal_input()
    if error:
        return error, 400
    document = build_catalog_proposal_document(offer, form)
    output = BytesIO()
    document.save(output)
    output.seek(0)
    safe_supplier = re.sub(r"[^0-9A-Za-zА-Яа-я_-]+", "_", offer["supplier"]).strip("_")
    return send_file(
        output,
        as_attachment=True,
        download_name=f"КП_полный_каталог_{safe_supplier}_{date.today().isoformat()}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@zbi_bp.route("/proposal.docx", methods=["POST"])
def proposal_docx():
    products, delivery_rows = load_catalog()
    product_by_id = {item["id"]: item for item in products}
    form = {
        "address": request.form.get("address", "").strip(),
        "markup": request.form.get("markup", "10"),
        "client_company": request.form.get("client_company", "").strip(),
        "client_inn": request.form.get("client_inn", "").strip(),
        "client_contact_name": request.form.get("client_contact_name", "").strip(),
        "client_contact": request.form.get("client_contact", "").strip(),
    }
    cart = parse_cart(request.form.get("items_json", "[]"), product_by_id)
    if not cart:
        return "Добавьте хотя бы одно изделие в заявку.", 400
    lat, lon, is_moscow = geocode(form["address"])
    if lat is None:
        return "Адрес не найден. Вернитесь в калькулятор и проверьте адрес.", 400
    markup = max(0, _float(form["markup"], 0))
    vehicle_overrides = parse_vehicle_overrides(request.form.get("vehicles_json", "{}"), delivery_rows)
    manual_loads = parse_manual_loads(request.form.get("loads_json", "{}"), product_by_id)
    quote = build_quote(
        cart, product_by_id, delivery_rows, lat, lon, markup, is_moscow,
        vehicle_overrides, manual_loads
    )
    if not quote["deliveries"]:
        return "Не удалось подобрать транспорт для коммерческого предложения.", 400
    document = build_proposal_document(quote, form, markup)
    output = BytesIO()
    document.save(output)
    output.seek(0)
    return send_file(
        output,
        as_attachment=True,
        download_name=f"КП_АР-ФАРВАТЕР_{date.today().isoformat()}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
